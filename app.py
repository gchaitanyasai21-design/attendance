import math
import os
import sqlite3
import secrets
import smtplib
import time
from datetime import date, timedelta
from email.message import EmailMessage
from typing import Dict, List, Optional
import xml.etree.ElementTree as ET
import zipfile
from flask import Flask, flash, g, redirect, render_template, request, session, url_for
from werkzeug.security import check_password_hash, generate_password_hash
try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOWNLOADS_DIR = os.path.join(os.environ.get("USERPROFILE", ""), "Downloads")
DB_PATH = os.path.join(BASE_DIR, "attendance.db")
MIN_PERCENTAGE = 75.0
COLLEGE_NAME = "MIC COLLEGE OF TECHNOLOGY"
SNAPSHOT_DATE = date(2026, 1, 31)
ATTENDANCE_WINDOW_START = date(2026, 1, 19)
ATTENDANCE_WINDOW_END = date(2026, 1, 31)
ATTENDANCE_TOTAL_CLASSES = 61
STUDENT_DATA_PATHS = [
    os.environ.get("STUDENT_DATA_XLSX", "").strip(),
    os.path.join(BASE_DIR, "25Batch_Students_data.xlsx"),
]
_env_class_docx_paths = [
    path.strip()
    for path in os.environ.get("CLASS_ATTENDANCE_DOCX", "").split(os.pathsep)
    if path.strip()
]
CLASS_DOCX_PATHS = _env_class_docx_paths + [
    os.path.join(DOWNLOADS_DIR, "AIDS-A  .docx"),
    os.path.join(DOWNLOADS_DIR, "AIDS-B  .docx"),
    os.path.join(DOWNLOADS_DIR, "AIDS-C  .docx"),
    os.path.join(DOWNLOADS_DIR, "AIML-A  .docx"),
    os.path.join(DOWNLOADS_DIR, "AIML-B  .docx"),
    os.path.join(DOWNLOADS_DIR, "CIVIL  ATTENDANCE.docx"),
    os.path.join(DOWNLOADS_DIR, "CSE-A .docx"),
    os.path.join(DOWNLOADS_DIR, "CSE-B .docx"),
    os.path.join(DOWNLOADS_DIR, "CSE-C.docx"),
    os.path.join(DOWNLOADS_DIR, "CSE-D .docx"),
    os.path.join(DOWNLOADS_DIR, "CSE-E  .docx"),
    os.path.join(DOWNLOADS_DIR, "CSE-F .docx"),
    os.path.join(DOWNLOADS_DIR, "ECE-A .docx"),
    os.path.join(DOWNLOADS_DIR, "ECE-B.docx"),
    os.path.join(DOWNLOADS_DIR, "ECE-C .docx"),
    os.path.join(DOWNLOADS_DIR, "EEE .docx"),
    os.path.join(DOWNLOADS_DIR, "IT-A .docx"),
    os.path.join(DOWNLOADS_DIR, "IT-B .docx"),
]
TIME_TABLE_DOCX_PATHS = [
    os.environ.get("TIME_TABLE_DOCX", "").strip(),
    os.path.join(BASE_DIR, "data", "time_table.docx"),
    os.path.join(BASE_DIR, "time_table.docx"),
    os.path.join(BASE_DIR, "time table .docx"),
]

CIVIL_ATTENDANCE_DATA = [
    ("25H71A0101", 55, 61),
    ("25H71A0102", 9, 61),
    ("25H71A0103", 0, 61),
    ("25H71A0104", 44, 61),
    ("25H71A0105", 0, 61),
    ("25H71A0106", 9, 61),
    ("25H71A0107", 14, 61),
    ("25H71A0108", 16, 61),
    ("25H71A0109", 18, 61),
    ("25H71A0110", 24, 61),
    ("25H71A0111", 34, 61),
    ("25H71A0112", 3, 61),
    ("25H71A0113", 55, 61),
    ("25H71A0114", 51, 61),
    ("25H71A0117", 17, 61),
    ("25H71A0118", 8, 61),
    ("25H71A0119", 43, 61),
    ("25H71A0120", 46, 61),
    ("25H71A0121", 57, 61),
    ("25H71A0122", 31, 61),
    ("25H71A0123", 5, 61),
    ("25H71A0124", 7, 61),
    ("25H71A0125", 30, 61),
    ("25H71A0126", 17, 61),
]

app = Flask(__name__)
app.config["SECRET_KEY"] = "change-this-secret-key"
app.config["SMTP_HOST"] = os.environ.get("SMTP_HOST", "")
app.config["SMTP_PORT"] = int(os.environ.get("SMTP_PORT", "587"))
app.config["SMTP_USER"] = os.environ.get("SMTP_USER", "")
app.config["SMTP_PASSWORD"] = os.environ.get("SMTP_PASSWORD", "")
app.config["SMTP_USE_TLS"] = os.environ.get("SMTP_USE_TLS", "1") == "1"
app.config["MAIL_FROM"] = os.environ.get("MAIL_FROM", app.config["SMTP_USER"])
OTP_EXPIRY_SECONDS = 600
DEFAULT_TEACHER_USERNAME = os.environ.get("TEACHER_USERNAME", "teacher").strip().lower()
DEFAULT_TEACHER_PASSWORD = os.environ.get("TEACHER_PASSWORD", "teacher123").strip()
DEFAULT_TEACHER_EMAIL = os.environ.get("TEACHER_EMAIL", "teacher@college.local").strip().lower()
ADMIN_TEACHER_USERNAME = os.environ.get("ADMIN_TEACHER_USERNAME", "HODMIC@").strip().lower()
ADMIN_TEACHER_PASSWORD = os.environ.get("ADMIN_TEACHER_PASSWORD", "hod123@").strip()
ADMIN_TEACHER_EMAIL = os.environ.get("ADMIN_TEACHER_EMAIL", "hodmic@college.local").strip().lower()

_TIMETABLE_CACHE: Dict[str, dict] = {}
_TIMETABLE_CACHE_PATH: Optional[str] = None
_TIMETABLE_CACHE_MTIME: Optional[float] = None


def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(DB_PATH, timeout=60)
        g.db.row_factory = sqlite3.Row
    return g.db


@app.teardown_appcontext
def close_db(exception):
    db = g.pop("db", None)
    if db is not None:
        db.close()


@app.context_processor
def inject_global_template_vars():
    return {"college_name": COLLEGE_NAME}


@app.before_request
def normalize_legacy_session():
    if "role" not in session and "student_id" in session:
        session["role"] = "student"
    if "role" not in session and "teacher_id" in session:
        session["role"] = "teacher"


def get_logged_in_role():
    return session.get("role")


def get_logged_in_student_id():
    if get_logged_in_role() == "student":
        return session.get("student_id")
    return None


def get_logged_in_teacher_id():
    if get_logged_in_role() == "teacher":
        return session.get("teacher_id")
    return None


def get_logged_in_teacher(db):
    teacher_id = get_logged_in_teacher_id()
    if not teacher_id:
        return None
    return db.execute("SELECT * FROM teachers WHERE id = ?", (teacher_id,)).fetchone()


def is_admin_teacher(db):
    teacher = get_logged_in_teacher(db)
    if not teacher:
        return False
    return teacher["username"].strip().lower() == ADMIN_TEACHER_USERNAME


def init_db():
    db = get_db()
    db.executescript(
        """
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            first_name TEXT,
            last_name TEXT,
            roll_no TEXT NOT NULL UNIQUE,
            email TEXT,
            password_hash TEXT,
            reset_otp TEXT,
            reset_otp_expires INTEGER,
            department TEXT,
            semester TEXT
        );

        CREATE TABLE IF NOT EXISTS attendance_records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id INTEGER NOT NULL,
            attendance_date TEXT NOT NULL,
            subject TEXT NOT NULL,
            status INTEGER NOT NULL CHECK(status IN (0, 1)),
            FOREIGN KEY(student_id) REFERENCES students(id)
        );

        CREATE TABLE IF NOT EXISTS teachers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            username TEXT NOT NULL UNIQUE,
            email TEXT,
            password_hash TEXT NOT NULL,
            reset_otp TEXT,
            reset_otp_expires INTEGER
        );
        """
    )
    _migrate_students_table(db)
    _migrate_teachers_table(db)
    _seed_default_teacher(db)
    _seed_admin_teacher(db)
    db.commit()


def _migrate_students_table(db):
    columns = {
        row["name"] for row in db.execute("PRAGMA table_info(students)").fetchall()
    }

    if "email" not in columns:
        db.execute("ALTER TABLE students ADD COLUMN email TEXT")
    if "first_name" not in columns:
        db.execute("ALTER TABLE students ADD COLUMN first_name TEXT")
    if "last_name" not in columns:
        db.execute("ALTER TABLE students ADD COLUMN last_name TEXT")
    if "password_hash" not in columns:
        db.execute("ALTER TABLE students ADD COLUMN password_hash TEXT")
    if "reset_otp" not in columns:
        db.execute("ALTER TABLE students ADD COLUMN reset_otp TEXT")
    if "reset_otp_expires" not in columns:
        db.execute("ALTER TABLE students ADD COLUMN reset_otp_expires INTEGER")
    _ensure_students_email_not_unique(db)

    db.execute(
        "CREATE UNIQUE INDEX IF NOT EXISTS idx_students_roll_no ON students(roll_no)"
    )
    email_index = db.execute(
        "SELECT sql FROM sqlite_master WHERE type = 'index' AND name = 'idx_students_email'"
    ).fetchone()
    email_index_sql = (email_index["sql"] or "").upper() if email_index else ""
    if "UNIQUE INDEX" in email_index_sql:
        db.execute("DROP INDEX idx_students_email")
    db.execute(
        "CREATE INDEX IF NOT EXISTS idx_students_email ON students(email)"
    )

    students_without_password = db.execute(
        "SELECT id, roll_no FROM students WHERE password_hash IS NULL OR password_hash = ''"
    ).fetchall()
    for row in students_without_password:
        db.execute(
            "UPDATE students SET password_hash = ? WHERE id = ?",
            (generate_password_hash(row["roll_no"]), row["id"]),
        )

    students_without_email = db.execute(
        "SELECT id, roll_no FROM students WHERE email IS NULL OR email = ''"
    ).fetchall()
    for row in students_without_email:
        generated_email = f"{row['roll_no'].lower()}@college.local"
        db.execute(
            "UPDATE students SET email = ? WHERE id = ?",
            (generated_email, row["id"]),
        )


def _ensure_students_email_not_unique(db):
    students_table_sql = db.execute(
        "SELECT sql FROM sqlite_master WHERE type = 'table' AND name = 'students'"
    ).fetchone()
    sql_text = (students_table_sql["sql"] or "").upper() if students_table_sql else ""
    if "EMAIL TEXT UNIQUE" not in sql_text:
        return

    db.execute("PRAGMA foreign_keys = OFF")
    db.executescript(
        """
        CREATE TABLE students_new (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            first_name TEXT,
            last_name TEXT,
            roll_no TEXT NOT NULL UNIQUE,
            email TEXT,
            password_hash TEXT,
            reset_otp TEXT,
            reset_otp_expires INTEGER,
            department TEXT,
            semester TEXT
        );

        INSERT INTO students_new(
            id, name, first_name, last_name, roll_no, email, password_hash,
            reset_otp, reset_otp_expires, department, semester
        )
        SELECT
            id, name, first_name, last_name, roll_no, email, password_hash,
            reset_otp, reset_otp_expires, department, semester
        FROM students;

        DROP TABLE students;
        ALTER TABLE students_new RENAME TO students;
        """
    )
    db.execute("PRAGMA foreign_keys = ON")


def _seed_default_teacher(db):
    if not DEFAULT_TEACHER_USERNAME or not DEFAULT_TEACHER_PASSWORD:
        return

    teacher = db.execute(
        "SELECT id FROM teachers WHERE username = ?",
        (DEFAULT_TEACHER_USERNAME,),
    ).fetchone()
    if teacher:
        return

    db.execute(
        """
        INSERT INTO teachers(name, username, password_hash)
        VALUES (?, ?, ?)
        """,
        (
            "Default Faculty",
            DEFAULT_TEACHER_USERNAME,
            generate_password_hash(DEFAULT_TEACHER_PASSWORD),
        ),
    )
    db.execute(
        "UPDATE teachers SET email = ? WHERE username = ?",
        (DEFAULT_TEACHER_EMAIL, DEFAULT_TEACHER_USERNAME),
    )


def _seed_admin_teacher(db):
    if not ADMIN_TEACHER_USERNAME or not ADMIN_TEACHER_PASSWORD:
        return

    teacher = db.execute(
        "SELECT id FROM teachers WHERE username = ?",
        (ADMIN_TEACHER_USERNAME,),
    ).fetchone()
    if teacher:
        return

    db.execute(
        """
        INSERT INTO teachers(name, username, password_hash)
        VALUES (?, ?, ?)
        """,
        (
            "Admin Faculty",
            ADMIN_TEACHER_USERNAME,
            generate_password_hash(ADMIN_TEACHER_PASSWORD),
        ),
    )
    db.execute(
        "UPDATE teachers SET email = ? WHERE username = ?",
        (ADMIN_TEACHER_EMAIL, ADMIN_TEACHER_USERNAME),
    )


def _migrate_teachers_table(db):
    columns = {
        row["name"] for row in db.execute("PRAGMA table_info(teachers)").fetchall()
    }
    if "email" not in columns:
        db.execute("ALTER TABLE teachers ADD COLUMN email TEXT")
    if "reset_otp" not in columns:
        db.execute("ALTER TABLE teachers ADD COLUMN reset_otp TEXT")
    if "reset_otp_expires" not in columns:
        db.execute("ALTER TABLE teachers ADD COLUMN reset_otp_expires INTEGER")

    db.execute(
        "CREATE UNIQUE INDEX IF NOT EXISTS idx_teachers_username ON teachers(username)"
    )
    db.execute(
        "CREATE INDEX IF NOT EXISTS idx_teachers_email ON teachers(email)"
    )

    teachers_without_email = db.execute(
        "SELECT id, username FROM teachers WHERE email IS NULL OR email = ''"
    ).fetchall()
    for row in teachers_without_email:
        generated_email = f"{row['username']}@college.local"
        db.execute(
            "UPDATE teachers SET email = ? WHERE id = ?",
            (generated_email, row["id"]),
        )


def calculate_stats(student_id: int):
    db = get_db()
    totals = db.execute(
        """
        SELECT
            COUNT(*) AS total_classes,
            COALESCE(SUM(status), 0) AS attended_classes
        FROM attendance_records
        WHERE student_id = ?
        """,
        (student_id,),
    ).fetchone()

    total = totals["total_classes"]
    attended = totals["attended_classes"]
    percentage = round((attended / total) * 100, 2) if total else 0.0

    required_attended_now = math.ceil((MIN_PERCENTAGE / 100) * total)
    shortage_now = max(0, required_attended_now - attended)

    classes_needed_to_recover = 0
    if percentage < MIN_PERCENTAGE and MIN_PERCENTAGE < 100:
        m = MIN_PERCENTAGE / 100
        classes_needed_to_recover = max(0, math.ceil((m * total - attended) / (1 - m)))

    return {
        "total": total,
        "attended": attended,
        "percentage": percentage,
        "shortage_now": shortage_now,
        "needed_to_recover": classes_needed_to_recover,
    }


def _build_spread_statuses(attended_classes: int, total_classes: int):
    statuses = []
    previous_target = 0
    for i in range(total_classes):
        current_target = round(((i + 1) * attended_classes) / total_classes)
        statuses.append(1 if current_target > previous_target else 0)
        previous_target = current_target
    return statuses


def _build_attendance_class_dates(total_classes: int):
    if total_classes <= 0:
        return []

    teaching_dates: List[date] = []
    current = ATTENDANCE_WINDOW_START
    while current <= ATTENDANCE_WINDOW_END:
        if current.weekday() != 6:  # Exclude Sunday.
            teaching_dates.append(current)
        current += timedelta(days=1)

    if not teaching_dates:
        teaching_dates = [ATTENDANCE_WINDOW_END]

    base = total_classes // len(teaching_dates)
    remainder = total_classes % len(teaching_dates)
    dates: List[str] = []
    for idx, dt in enumerate(teaching_dates):
        classes_for_day = base + (1 if idx < remainder else 0)
        for _ in range(classes_for_day):
            dates.append(dt.isoformat())
    return dates


def _roman_to_int(value: str):
    mapping = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6}
    return mapping.get(value.strip().upper())


def _normalize_section_label(value: str) -> str:
    text = " ".join((value or "").split())
    text = text.replace(" - ", "-").replace("- ", "-").replace(" -", "-")
    return text.upper()


def _clean_cell_text(value: str) -> str:
    if not value:
        return ""
    text = value.replace("\r", "\n")
    lines = []
    for line in text.split("\n"):
        cleaned = " ".join(line.split())
        if cleaned:
            lines.append(cleaned)
    return "\n".join(lines)


def _looks_like_schedule_table(table) -> bool:
    if not table.rows or len(table.rows) < 3:
        return False
    row1 = (table.rows[1].cells[0].text or "").upper()
    row2 = (table.rows[2].cells[0].text or "").upper()
    return "PERIOD" in row1 and "DAY" in row2


def _looks_like_course_table(table) -> bool:
    if not table.rows or len(table.rows) < 2:
        return False
    header_text = " ".join((cell.text or "").upper() for cell in table.rows[1].cells)
    return "COURSE CODE" in header_text and "COURSE TITLE" in header_text


def _extract_table_rows(table) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in table.rows:
        rows.append([_clean_cell_text(cell.text) for cell in row.cells])
    return rows


def _parse_timetable_docx(docx_path: str) -> Dict[str, dict]:
    if Document is None:
        return {}
    if not os.path.exists(docx_path):
        return {}

    doc = Document(docx_path)
    results: Dict[str, dict] = {}
    idx = 0

    while idx < len(doc.tables):
        table = doc.tables[idx]
        if _looks_like_schedule_table(table):
            class_label_raw = (table.rows[0].cells[0].text or "").strip()
            class_label = _normalize_section_label(class_label_raw)
            schedule_rows = _extract_table_rows(table)
            courses = None

            if idx + 1 < len(doc.tables) and _looks_like_course_table(doc.tables[idx + 1]):
                course_rows = _extract_table_rows(doc.tables[idx + 1])
                if course_rows and "CLASS IN-CHARGE" in (course_rows[0][0] or "").upper():
                    course_rows = course_rows[1:]
                if course_rows:
                    courses = {
                        "headers": course_rows[0],
                        "rows": course_rows[1:],
                    }
                idx += 1

            if class_label:
                results[class_label] = {
                    "label": class_label_raw or class_label,
                    "schedule": schedule_rows,
                    "courses": courses,
                }
        idx += 1

    return results


def _dedupe_schedule_period_columns(schedule: List[List[str]]) -> List[List[str]]:
    if not schedule or len(schedule) < 2:
        return schedule

    period_row = schedule[1]
    if not period_row:
        return schedule

    keep_indexes = [0]
    seen_period_numbers = set()

    for idx in range(1, len(period_row)):
        label = (period_row[idx] or "").strip()
        if label.isdigit():
            if label in seen_period_numbers:
                continue
            seen_period_numbers.add(label)
        keep_indexes.append(idx)

    if len(keep_indexes) == len(period_row):
        return schedule

    deduped: List[List[str]] = []
    for row in schedule:
        deduped.append([row[i] if i < len(row) else "" for i in keep_indexes])
    return deduped


def _load_timetables() -> Dict[str, dict]:
    global _TIMETABLE_CACHE, _TIMETABLE_CACHE_PATH, _TIMETABLE_CACHE_MTIME

    for path in TIME_TABLE_DOCX_PATHS:
        if not path:
            continue
        if not os.path.exists(path):
            continue
        mtime = os.path.getmtime(path)
        if _TIMETABLE_CACHE and _TIMETABLE_CACHE_PATH == path and _TIMETABLE_CACHE_MTIME == mtime:
            return _TIMETABLE_CACHE
        data = _parse_timetable_docx(path)
        _TIMETABLE_CACHE = data
        _TIMETABLE_CACHE_PATH = path
        _TIMETABLE_CACHE_MTIME = mtime
        return data
    return {}


def _get_student_timetable(student_department: str) -> Optional[dict]:
    section_key = _normalize_section_label(student_department or "")
    if not section_key:
        return None
    timetables = _load_timetables()
    timetable = timetables.get(section_key)
    if not timetable:
        return None

    schedule = timetable.get("schedule") or []
    if len(schedule) >= 3:
        header_row = schedule[0]
        period_row = schedule[1]
        day_row = schedule[2]
        if (
            header_row
            and period_row
            and day_row
            and "PERIOD" in (period_row[0] or "").upper()
            and "DAY" in (day_row[0] or "").upper()
        ):
            adjusted = [row[:] for row in schedule]
            label = timetable.get("label") or header_row[0]
            new_header = [""] * len(header_row)
            new_header[0] = label
            adjusted[0] = new_header
            adjusted[1] = [period_row[0]] + header_row[1:]
            schedule = adjusted

    deduped_schedule = _dedupe_schedule_period_columns(schedule)
    if deduped_schedule is not schedule:
        timetable = {**timetable, "schedule": deduped_schedule}

    return timetable


def _parse_class_docx(docx_path: str):
    if Document is None:
        return None
    if not os.path.exists(docx_path):
        return None

    doc = Document(docx_path)
    class_label = None
    semester_value = None

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        if text.upper().startswith("CLASS:"):
            class_text = text.split("CLASS:", 1)[1]
            class_text = class_text.split("Academic Year", 1)[0]
            class_text = " ".join(class_text.split())
            class_label = class_text.replace(" - ", "-").replace("- ", "-").replace(" -", "-")
        if "SEMESTER" in text.upper():
            marker = "B.Tech-"
            if marker in text:
                roman = text.split(marker, 1)[1].split(" ", 1)[0]
                semester_value = _roman_to_int(roman) or semester_value

    if not doc.tables:
        return None

    table = doc.tables[0]
    if not table.rows:
        return None

    header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
    try:
        idx_roll = header_cells.index("regd. no.")
    except ValueError:
        idx_roll = None
    try:
        idx_name = header_cells.index("name of the student")
    except ValueError:
        idx_name = None
    try:
        idx_attended = header_cells.index("classes attended")
    except ValueError:
        idx_attended = None
    try:
        idx_total = header_cells.index("classes conducted")
    except ValueError:
        idx_total = None

    if None in {idx_roll, idx_name, idx_attended, idx_total}:
        return None

    rows = []
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) <= max(idx_roll, idx_name, idx_attended, idx_total):
            continue
        roll_no = cells[idx_roll].strip()
        name = cells[idx_name].strip()
        attended = cells[idx_attended].strip()
        total = cells[idx_total].strip()
        if not roll_no:
            continue
        try:
            attended_val = int(float(attended))
            total_val = int(float(total))
        except ValueError:
            continue
        rows.append(
            {
                "roll_no": roll_no,
                "name": name or roll_no,
                "attended": attended_val,
                "total": total_val,
            }
        )

    return {
        "class_label": class_label or "",
        "semester": str(semester_value) if semester_value else "",
        "rows": rows,
    }


def import_class_docx_attendance():
    if Document is None:
        return

    db = get_db()
    for path in CLASS_DOCX_PATHS:
        parsed = _parse_class_docx(path)
        if not parsed or not parsed["rows"]:
            continue

        class_label = parsed["class_label"]
        semester = parsed["semester"]

        for row in parsed["rows"]:
            roll_no = row["roll_no"]
            name = row["name"]
            parts = [part for part in name.split() if part]
            first_name = parts[0] if parts else ""
            last_name = " ".join(parts[1:]) if len(parts) > 1 else ""

            student = db.execute(
                "SELECT id, email FROM students WHERE roll_no = ?",
                (roll_no,),
            ).fetchone()

            if student:
                db.execute(
                    """
                    UPDATE students
                    SET name = ?, first_name = ?, last_name = ?, department = ?, semester = ?
                    WHERE roll_no = ?
                    """,
                    (name, first_name, last_name, class_label, semester, roll_no),
                )
                student_id = student["id"]
            else:
                email = f"{roll_no.lower()}@college.local"
                db.execute(
                    """
                    INSERT INTO students(name, first_name, last_name, roll_no, email, password_hash, department, semester)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        name,
                        first_name,
                        last_name,
                        roll_no,
                        email,
                        generate_password_hash(roll_no),
                        class_label,
                        semester,
                    ),
                )
                student_id = db.execute(
                    "SELECT id FROM students WHERE roll_no = ?",
                    (roll_no,),
                ).fetchone()["id"]

            total_classes = ATTENDANCE_TOTAL_CLASSES
            attended_classes = min(max(0, row["attended"]), total_classes)

            db.execute("DELETE FROM attendance_records WHERE student_id = ?", (student_id,))
            if total_classes:
                statuses = _build_spread_statuses(attended_classes, total_classes)
                attendance_dates = _build_attendance_class_dates(total_classes)
                for attendance_date, status in zip(attendance_dates, statuses):
                    db.execute(
                        """
                        INSERT INTO attendance_records(student_id, attendance_date, subject, status)
                        VALUES (?, ?, ?, ?)
                        """,
                        (student_id, attendance_date, "Overall", status),
                    )

    db.commit()


def _read_students_from_xlsx(xlsx_path: str):
    ns = {
        "m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }

    with zipfile.ZipFile(xlsx_path) as zf:
        shared_strings = []
        if "xl/sharedStrings.xml" in zf.namelist():
            shared_root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
            for item in shared_root.findall("m:si", ns):
                texts = [node.text or "" for node in item.findall(".//m:t", ns)]
                shared_strings.append("".join(texts))

        workbook = ET.fromstring(zf.read("xl/workbook.xml"))
        rels = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
        rel_id_to_target = {
            rel.attrib["Id"]: rel.attrib["Target"]
            for rel in rels.findall(
                "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
            )
        }
        first_sheet = workbook.find("m:sheets/m:sheet", ns)
        if first_sheet is None:
            return []

        rel_id = first_sheet.attrib.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        )
        target = rel_id_to_target.get(rel_id, "")
        if not target:
            return []
        if not target.startswith("xl/"):
            target = f"xl/{target}"

        sheet_root = ET.fromstring(zf.read(target))
        rows = []
        for row in sheet_root.findall("m:sheetData/m:row", ns):
            values = []
            for cell in row.findall("m:c", ns):
                cell_type = cell.attrib.get("t")
                value_node = cell.find("m:v", ns)
                if value_node is None:
                    values.append("")
                    continue

                raw_value = value_node.text or ""
                if cell_type == "s" and raw_value.isdigit():
                    values.append(shared_strings[int(raw_value)])
                else:
                    values.append(raw_value)
            rows.append(values)

    if not rows:
        return []

    students = []
    for row in rows[1:]:
        if len(row) < 5:
            continue
        roll_no = row[0].strip()
        first_name = row[1].strip()
        last_name = row[2].strip()
        email = row[3].strip().lower()
        department = row[4].strip()
        if not roll_no or not email:
            continue
        full_name = " ".join(part for part in [first_name, last_name] if part).strip()
        students.append(
            {
                "roll_no": roll_no,
                "first_name": first_name,
                "last_name": last_name,
                "name": full_name or roll_no,
                "email": email,
                "department": department,
            }
        )

    return sorted(students, key=lambda row: row["roll_no"])


def import_students_data():
    xlsx_path = next(
        (path for path in STUDENT_DATA_PATHS if path and os.path.exists(path)),
        "",
    )
    if not xlsx_path:
        return

    db = get_db()
    students = _read_students_from_xlsx(xlsx_path)
    existing_passwords = {
        row["roll_no"]: row["password_hash"]
        for row in db.execute("SELECT roll_no, password_hash FROM students").fetchall()
    }
    for student in students:
        password_hash = existing_passwords.get(student["roll_no"]) or generate_password_hash(
            student["roll_no"]
        )
        db.execute(
            """
            INSERT INTO students(name, first_name, last_name, roll_no, email, password_hash, department, semester)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(roll_no) DO UPDATE SET
                name = excluded.name,
                first_name = excluded.first_name,
                last_name = excluded.last_name,
                email = excluded.email,
                department = excluded.department
            """,
            (
                student["name"],
                student["first_name"],
                student["last_name"],
                student["roll_no"],
                student["email"],
                password_hash,
                student["department"] or "N/A",
                "2",
            ),
        )

    db.execute(
        """
        DELETE FROM attendance_records
        WHERE student_id IN (
            SELECT id FROM students WHERE roll_no = 'CSE001'
        )
        """
    )
    db.execute("DELETE FROM students WHERE roll_no = 'CSE001'")
    db.commit()


def seed_civil_attendance_data():
    db = get_db()
    for roll_no, attended_classes, total_classes in sorted(
        CIVIL_ATTENDANCE_DATA, key=lambda row: row[0]
    ):
        student = db.execute(
            "SELECT id FROM students WHERE roll_no = ?",
            (roll_no,),
        ).fetchone()
        if not student:
            db.execute(
                """
                INSERT INTO students(name, first_name, last_name, roll_no, email, password_hash, department, semester)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    roll_no,
                    "",
                    "",
                    roll_no,
                    f"{roll_no.lower()}@mictech.edu.in",
                    generate_password_hash(roll_no),
                    "CIVIL",
                    "2",
                ),
            )
            student = db.execute(
                "SELECT id FROM students WHERE roll_no = ?",
                (roll_no,),
            ).fetchone()
        student_id = student["id"]

        db.execute("DELETE FROM attendance_records WHERE student_id = ?", (student_id,))
        total_classes = ATTENDANCE_TOTAL_CLASSES
        attended_classes = min(max(0, attended_classes), total_classes)
        statuses = _build_spread_statuses(attended_classes, total_classes)
        attendance_dates = _build_attendance_class_dates(total_classes)
        for attendance_date, status in zip(attendance_dates, statuses):
            db.execute(
                """
                INSERT INTO attendance_records(student_id, attendance_date, subject, status)
                VALUES (?, ?, ?, ?)
                """,
                (student_id, attendance_date, "Overall", status),
            )

    db.commit()


def send_otp_email(to_email: str, otp: str):
    if not app.config["SMTP_HOST"] or not app.config["SMTP_USER"] or not app.config["SMTP_PASSWORD"]:
        raise RuntimeError("SMTP is not configured.")

    message = EmailMessage()
    message["Subject"] = "Attendance App Password Reset OTP"
    message["From"] = app.config["MAIL_FROM"]
    message["To"] = to_email
    message.set_content(
        f"Your OTP for password reset is: {otp}\n"
        f"It will expire in {OTP_EXPIRY_SECONDS // 60} minutes."
    )

    with smtplib.SMTP(app.config["SMTP_HOST"], app.config["SMTP_PORT"]) as server:
        if app.config["SMTP_USE_TLS"]:
            server.starttls()
        server.login(app.config["SMTP_USER"], app.config["SMTP_PASSWORD"])
        server.send_message(message)


@app.route("/")
def home():
    role = get_logged_in_role()
    if role == "student" and get_logged_in_student_id():
        return redirect(url_for("dashboard"))
    if role == "teacher" and get_logged_in_teacher_id():
        return redirect(url_for("admin"))
    return redirect(url_for("login"))


@app.route("/login", methods=["GET", "POST"])
def login():
    role = get_logged_in_role()
    if role == "student" and get_logged_in_student_id():
        return redirect(url_for("dashboard"))
    if role == "teacher" and get_logged_in_teacher_id():
        return redirect(url_for("admin"))

    if request.method == "POST":
        login_type = request.form.get("login_type", "student").strip().lower()
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()

        if login_type not in {"student", "teacher"}:
            login_type = "student"

        if not username or not password:
            flash("Please enter username and password.", "error")
            return render_template("login.html", selected_login_type=login_type)

        db = get_db()
        if login_type == "student":
            student = db.execute(
                "SELECT * FROM students WHERE roll_no = ?",
                (username,),
            ).fetchone()

            if not student:
                flash("Student not found. Ask faculty to add your profile.", "error")
                return render_template("login.html", selected_login_type=login_type)

            password_hash = student["password_hash"]
            is_valid = bool(password_hash) and check_password_hash(password_hash, password)
            if not is_valid:
                flash("Invalid credentials.", "error")
                return render_template("login.html", selected_login_type=login_type)

            session.clear()
            session["role"] = "student"
            session["student_id"] = student["id"]
            return redirect(url_for("dashboard"))

        teacher = db.execute(
            "SELECT * FROM teachers WHERE username = ?",
            (username.lower(),),
        ).fetchone()
        if not teacher:
            flash("Faculty account not found.", "error")
            return render_template("login.html", selected_login_type=login_type)

        if not check_password_hash(teacher["password_hash"], password):
            flash("Invalid credentials.", "error")
            return render_template("login.html", selected_login_type=login_type)

        session.clear()
        session["role"] = "teacher"
        session["teacher_id"] = teacher["id"]
        return redirect(url_for("admin"))

    return render_template("login.html", selected_login_type="student")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if get_logged_in_role() == "teacher":
        return redirect(url_for("admin"))

    if request.method == "POST":
        roll_no = request.form.get("roll_no", "").strip()
        email = request.form.get("email", "").strip().lower()
        db = get_db()
        student = db.execute(
            "SELECT id, roll_no, email FROM students WHERE roll_no = ?",
            (roll_no,),
        ).fetchone()

        if not student or not student["email"] or student["email"].lower() != email:
            flash("Roll number and email do not match.", "error")
            return render_template("forgot_password.html")

        otp = f"{secrets.randbelow(1000000):06d}"
        expires_at = int(time.time()) + OTP_EXPIRY_SECONDS
        db.execute(
            "UPDATE students SET reset_otp = ?, reset_otp_expires = ? WHERE id = ?",
            (otp, expires_at, student["id"]),
        )
        db.commit()

        try:
            send_otp_email(email, otp)
            flash("OTP sent to your email.", "success")
        except Exception:
            flash(
                f"SMTP is not configured yet. Dev OTP for {roll_no}: {otp}",
                "error",
            )

        session["password_reset_student_id"] = student["id"]
        return redirect(url_for("verify_otp"))

    return render_template("forgot_password.html")


@app.route("/teacher/forgot-password", methods=["GET", "POST"])
def teacher_forgot_password():
    if get_logged_in_role() == "student":
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        email = request.form.get("email", "").strip().lower()
        db = get_db()
        teacher = db.execute(
            "SELECT id, username, email FROM teachers WHERE username = ?",
            (username,),
        ).fetchone()

        if not teacher or not teacher["email"] or teacher["email"].lower() != email:
            flash("Faculty username and email do not match.", "error")
            return render_template("teacher_forgot_password.html")

        otp = f"{secrets.randbelow(1000000):06d}"
        expires_at = int(time.time()) + OTP_EXPIRY_SECONDS
        db.execute(
            "UPDATE teachers SET reset_otp = ?, reset_otp_expires = ? WHERE id = ?",
            (otp, expires_at, teacher["id"]),
        )
        db.commit()

        try:
            send_otp_email(email, otp)
            flash("OTP sent to your email.", "success")
        except Exception:
            flash(
                f"SMTP is not configured yet. Dev OTP for {username}: {otp}",
                "error",
            )

        session["password_reset_teacher_id"] = teacher["id"]
        return redirect(url_for("teacher_verify_otp"))

    return render_template("teacher_forgot_password.html")


@app.route("/verify-otp", methods=["GET", "POST"])
def verify_otp():
    student_id = session.get("password_reset_student_id")
    if not student_id:
        flash("Start from forgot password first.", "error")
        return redirect(url_for("forgot_password"))

    if request.method == "POST":
        otp = request.form.get("otp", "").strip()
        new_password = request.form.get("new_password", "").strip()
        confirm_password = request.form.get("confirm_password", "").strip()

        if not otp or not new_password or not confirm_password:
            flash("All fields are required.", "error")
            return render_template("verify_otp.html")
        if new_password != confirm_password:
            flash("New password and confirm password do not match.", "error")
            return render_template("verify_otp.html")

        db = get_db()
        student = db.execute(
            "SELECT id, reset_otp, reset_otp_expires FROM students WHERE id = ?",
            (student_id,),
        ).fetchone()
        now_ts = int(time.time())
        if (
            not student
            or not student["reset_otp"]
            or student["reset_otp"] != otp
            or not student["reset_otp_expires"]
            or now_ts > student["reset_otp_expires"]
        ):
            flash("Invalid or expired OTP.", "error")
            return render_template("verify_otp.html")

        db.execute(
            """
            UPDATE students
            SET password_hash = ?, reset_otp = NULL, reset_otp_expires = NULL
            WHERE id = ?
            """,
            (generate_password_hash(new_password), student_id),
        )
        db.commit()
        session.pop("password_reset_student_id", None)
        flash("Password reset successful. Please login.", "success")
        return redirect(url_for("login"))

    return render_template("verify_otp.html")


@app.route("/teacher/verify-otp", methods=["GET", "POST"])
def teacher_verify_otp():
    teacher_id = session.get("password_reset_teacher_id")
    if not teacher_id:
        flash("Start from faculty forgot password first.", "error")
        return redirect(url_for("teacher_forgot_password"))

    if request.method == "POST":
        otp = request.form.get("otp", "").strip()
        new_password = request.form.get("new_password", "").strip()
        confirm_password = request.form.get("confirm_password", "").strip()

        if not otp or not new_password or not confirm_password:
            flash("All fields are required.", "error")
            return render_template("teacher_verify_otp.html")
        if new_password != confirm_password:
            flash("New password and confirm password do not match.", "error")
            return render_template("teacher_verify_otp.html")

        db = get_db()
        teacher = db.execute(
            "SELECT id, reset_otp, reset_otp_expires FROM teachers WHERE id = ?",
            (teacher_id,),
        ).fetchone()
        now_ts = int(time.time())
        if (
            not teacher
            or not teacher["reset_otp"]
            or teacher["reset_otp"] != otp
            or not teacher["reset_otp_expires"]
            or now_ts > teacher["reset_otp_expires"]
        ):
            flash("Invalid or expired OTP.", "error")
            return render_template("teacher_verify_otp.html")

        db.execute(
            """
            UPDATE teachers
            SET password_hash = ?, reset_otp = NULL, reset_otp_expires = NULL
            WHERE id = ?
            """,
            (generate_password_hash(new_password), teacher_id),
        )
        db.commit()
        session.pop("password_reset_teacher_id", None)
        flash("Faculty password reset successful. Please login.", "success")
        return redirect(url_for("login"))

    return render_template("teacher_verify_otp.html")


@app.route("/change-password", methods=["GET", "POST"])
def change_password():
    student_id = get_logged_in_student_id()
    if not student_id:
        return redirect(url_for("login"))

    if request.method == "POST":
        current_password = request.form.get("current_password", "").strip()
        new_password = request.form.get("new_password", "").strip()
        confirm_password = request.form.get("confirm_password", "").strip()

        if not current_password or not new_password or not confirm_password:
            flash("All fields are required.", "error")
            return render_template("change_password.html")
        if new_password != confirm_password:
            flash("New password and confirm password do not match.", "error")
            return render_template("change_password.html")

        db = get_db()
        student = db.execute(
            "SELECT id, password_hash FROM students WHERE id = ?",
            (student_id,),
        ).fetchone()
        if not student or not check_password_hash(student["password_hash"], current_password):
            flash("Current password is incorrect.", "error")
            return render_template("change_password.html")

        db.execute(
            "UPDATE students SET password_hash = ? WHERE id = ?",
            (generate_password_hash(new_password), student_id),
        )
        db.commit()
        flash("Password updated successfully.", "success")
        return redirect(url_for("dashboard"))

    return render_template("change_password.html")


@app.route("/teacher/change-credentials", methods=["GET", "POST"])
def teacher_change_credentials():
    teacher_id = get_logged_in_teacher_id()
    if not teacher_id:
        return redirect(url_for("login"))

    db = get_db()
    teacher = db.execute(
        "SELECT id, username, email, password_hash FROM teachers WHERE id = ?",
        (teacher_id,),
    ).fetchone()
    if not teacher:
        session.clear()
        return redirect(url_for("login"))

    if request.method == "POST":
        current_password = request.form.get("current_password", "").strip()
        new_username = request.form.get("new_username", "").strip().lower()
        new_email = request.form.get("new_email", "").strip().lower()
        new_password = request.form.get("new_password", "").strip()
        confirm_password = request.form.get("confirm_password", "").strip()

        if not current_password:
            flash("Current password is required.", "error")
            return render_template("teacher_change_credentials.html", teacher=teacher)
        if not check_password_hash(teacher["password_hash"], current_password):
            flash("Current password is incorrect.", "error")
            return render_template("teacher_change_credentials.html", teacher=teacher)
        if not new_username:
            flash("Username is required.", "error")
            return render_template("teacher_change_credentials.html", teacher=teacher)
        if not new_email:
            flash("Email is required.", "error")
            return render_template("teacher_change_credentials.html", teacher=teacher)
        if (new_password or confirm_password) and new_password != confirm_password:
            flash("New password and confirm password do not match.", "error")
            return render_template("teacher_change_credentials.html", teacher=teacher)

        next_password_hash = teacher["password_hash"]
        if new_password:
            next_password_hash = generate_password_hash(new_password)

        try:
            db.execute(
                """
                UPDATE teachers
                SET username = ?, email = ?, password_hash = ?
                WHERE id = ?
                """,
                (new_username, new_email, next_password_hash, teacher_id),
            )
            db.commit()
        except sqlite3.IntegrityError:
            flash("Username already exists.", "error")
            return render_template("teacher_change_credentials.html", teacher=teacher)

        flash("Faculty credentials updated successfully.", "success")
        return redirect(url_for("admin"))

    return render_template("teacher_change_credentials.html", teacher=teacher)


@app.route("/change-email", methods=["GET", "POST"])
def change_email():
    student_id = get_logged_in_student_id()
    if not student_id:
        return redirect(url_for("login"))

    flash("Email changes are disabled. Use your college main email ID.", "error")
    return redirect(url_for("dashboard"))


@app.route("/dashboard")
def dashboard():
    student_id = get_logged_in_student_id()
    if not student_id:
        return redirect(url_for("login"))

    db = get_db()
    student = db.execute("SELECT * FROM students WHERE id = ?", (student_id,)).fetchone()
    stats = calculate_stats(student_id)
    timetable = _get_student_timetable(student["department"])

    records = db.execute(
        """
        SELECT attendance_date, subject, status
        FROM attendance_records
        WHERE student_id = ?
        ORDER BY attendance_date DESC, subject ASC, id DESC
        LIMIT 25
        """,
        (student_id,),
    ).fetchall()

    return render_template(
        "dashboard.html",
        student=student,
        stats=stats,
        min_percentage=MIN_PERCENTAGE,
        records=records,
        timetable=timetable,
    )
@app.route("/timetable")
def timetable():
    student_id = get_logged_in_student_id()
    if not student_id:
        return redirect(url_for("login"))

    db = get_db()
    student = db.execute("SELECT * FROM students WHERE id = ?", (student_id,)).fetchone()
    timetable = _get_student_timetable(student["department"])

    return render_template(
        "timetable.html",
        student=student,
        timetable=timetable,
    )


@app.route("/admin", methods=["GET", "POST"])
def admin():
    if not get_logged_in_teacher_id():
        return redirect(url_for("login"))

    db = get_db()
    current_teacher = get_logged_in_teacher(db)
    can_manage_teachers = is_admin_teacher(db)
    attendance_roll_no = request.args.get("attendance_roll_no", "").strip()
    attendance_student = None
    attendance_records = []
    attendance_stats = None
    attendance_percentage = 0.0


    if request.method == "POST":
        action = request.form.get("action")
        redirect_roll_no = ""

        if action == "add_attendance":
            roll_no = request.form.get("roll_no", "").strip()
            subject = request.form.get("subject", "").strip()
            status = request.form.get("status", "0").strip()
            attendance_date = request.form.get("attendance_date", "").strip() or date.today().isoformat()

            student = db.execute(
                "SELECT id FROM students WHERE roll_no = ?", (roll_no,)
            ).fetchone()
            if not student or not subject:
                flash("Valid roll number and subject are required.", "error")
            else:
                db.execute(
                    "INSERT INTO attendance_records(student_id, attendance_date, subject, status) VALUES (?, ?, ?, ?)",
                    (student["id"], attendance_date, subject, 1 if status == "1" else 0),
                )
                db.commit()
                flash("Attendance record added.", "success")
                redirect_roll_no = roll_no

        if action == "add_teacher":
            if not can_manage_teachers:
                flash("Only the admin user can manage faculty.", "error")
                return redirect(url_for("admin"))

            name = request.form.get("name", "").strip()
            username = request.form.get("username", "").strip().lower()
            email = request.form.get("email", "").strip().lower()
            password = request.form.get("password", "").strip()

            if not name or not username or not email or not password:
                flash("Name, username, email, and password are required.", "error")
            else:
                try:
                    db.execute(
                        """
                        INSERT INTO teachers(name, username, email, password_hash)
                        VALUES (?, ?, ?, ?)
                        """,
                        (name, username, email, generate_password_hash(password)),
                    )
                    db.commit()
                    flash("Faculty added.", "success")
                except sqlite3.IntegrityError:
                    flash("Faculty username already exists.", "error")

        if action == "remove_teacher":
            if not can_manage_teachers:
                flash("Only the admin user can manage faculty.", "error")
                return redirect(url_for("admin"))

            username = request.form.get("username", "").strip().lower()
            if not username:
                flash("Faculty username is required.", "error")
            elif current_teacher and username == current_teacher["username"].strip().lower():
                flash("You cannot remove the current logged-in faculty.", "error")
            elif username == ADMIN_TEACHER_USERNAME:
                flash("You cannot remove the admin faculty.", "error")
            else:
                deleted = db.execute(
                    "DELETE FROM teachers WHERE username = ?",
                    (username,),
                ).rowcount
                db.commit()
                if deleted:
                    flash("Faculty removed.", "success")
                else:
                    flash("Faculty not found.", "error")

        if action in {"update_attendance", "delete_attendance", "update_attendance_bulk"}:
            roll_no = request.form.get("attendance_roll_no", "").strip()
            redirect_roll_no = roll_no or redirect_roll_no
            if action in {"update_attendance", "delete_attendance"}:
                record_id = request.form.get("record_id", "").strip()
                if not record_id.isdigit():
                    flash("Invalid attendance record.", "error")
                    return redirect(
                        url_for("admin", attendance_roll_no=redirect_roll_no)
                        if redirect_roll_no
                        else url_for("admin")
                    )

            if action == "update_attendance":
                status = request.form.get("status", "0").strip()
                db.execute(
                    "UPDATE attendance_records SET status = ? WHERE id = ?",
                    (1 if status == "1" else 0, int(record_id)),
                )
                db.commit()
                flash("Attendance updated.", "success")

            if action == "delete_attendance":
                db.execute("DELETE FROM attendance_records WHERE id = ?", (int(record_id),))
                db.commit()
                flash("Attendance deleted.", "success")

            if action == "update_attendance_bulk":
                record_ids = []
                for raw_id in request.form.getlist("record_ids"):
                    raw_id = (raw_id or "").strip()
                    if raw_id.isdigit():
                        record_ids.append(int(raw_id))
                record_ids = list(dict.fromkeys(record_ids))

                if not record_ids:
                    flash("No attendance records selected for update.", "error")
                else:
                    updates = []
                    for rec_id in record_ids:
                        status = request.form.get(f"status_{rec_id}", "0").strip()
                        updates.append((1 if status == "1" else 0, rec_id))
                    db.executemany(
                        "UPDATE attendance_records SET status = ? WHERE id = ?",
                        updates,
                    )
                    db.commit()
                    flash(f"Saved {len(updates)} attendance updates.", "success")

        return redirect(
            url_for("admin", attendance_roll_no=redirect_roll_no)
            if redirect_roll_no
            else url_for("admin")
        )

    if attendance_roll_no:
        attendance_student = db.execute(
            "SELECT id, name, roll_no FROM students WHERE roll_no = ?",
            (attendance_roll_no,),
        ).fetchone()
        if attendance_student:
            attendance_records = db.execute(
                """
                SELECT id, attendance_date, subject, status
                FROM attendance_records
                WHERE student_id = ?
                ORDER BY attendance_date DESC, subject ASC, id DESC
                """,
                (attendance_student["id"],),
            ).fetchall()
            attendance_stats = calculate_stats(attendance_student["id"])
            attendance_percentage = attendance_stats["percentage"]
        else:
            flash("Student not found for attendance lookup.", "error")
    return render_template(
        "admin.html",
        can_manage_teachers=can_manage_teachers,
        attendance_roll_no=attendance_roll_no,
        attendance_student=attendance_student,
        attendance_records=attendance_records,
        attendance_stats=attendance_stats,
        attendance_percentage=attendance_percentage,
    )


@app.route("/teacher/students")
def teacher_students():
    if not get_logged_in_teacher_id():
        return redirect(url_for("login"))

    db = get_db()
    search_query = request.args.get("q", "").strip()
    students = []
    if search_query:
        wildcard = f"%{search_query.lower()}%"
        students = db.execute(
            """
            SELECT id, name, first_name, last_name, roll_no, email, department, semester
            FROM students
            WHERE
                LOWER(COALESCE(name, '')) LIKE ?
                OR LOWER(COALESCE(first_name, '')) LIKE ?
                OR LOWER(COALESCE(last_name, '')) LIKE ?
                OR LOWER(COALESCE(roll_no, '')) LIKE ?
                OR LOWER(COALESCE(email, '')) LIKE ?
                OR LOWER(COALESCE(department, '')) LIKE ?
                OR LOWER(COALESCE(semester, '')) LIKE ?
            ORDER BY roll_no
            """,
            (
                wildcard,
                wildcard,
                wildcard,
                wildcard,
                wildcard,
                wildcard,
                wildcard,
            ),
        ).fetchall()
    else:
        students = db.execute(
            """
            SELECT id, name, first_name, last_name, roll_no, email, department, semester
            FROM students
            ORDER BY roll_no
            """
        ).fetchall()

    return render_template(
        "teacher_students.html",
        students=students,
        search_query=search_query,
    )


@app.route("/teacher/notifications")
def teacher_notifications():
    if not get_logged_in_teacher_id():
        return redirect(url_for("login"))

    db = get_db()
    low_attendance_students = db.execute(
        """
        SELECT
            s.id,
            s.name,
            s.roll_no,
            s.department,
            s.semester,
            SUM(CASE WHEN ar.status = 1 THEN 1 ELSE 0 END) AS attended_classes,
            COUNT(ar.id) AS total_classes,
            ROUND(
                (100.0 * SUM(CASE WHEN ar.status = 1 THEN 1 ELSE 0 END)) / COUNT(ar.id),
                2
            ) AS attendance_percentage
        FROM students s
        LEFT JOIN attendance_records ar ON ar.student_id = s.id
        GROUP BY s.id, s.name, s.roll_no, s.department, s.semester
        HAVING COUNT(ar.id) > 0
           AND ((100.0 * SUM(CASE WHEN ar.status = 1 THEN 1 ELSE 0 END)) / COUNT(ar.id)) < 30
        ORDER BY attendance_percentage ASC, s.roll_no ASC
        """
    ).fetchall()

    return render_template(
        "teacher_notifications.html",
        low_attendance_students=low_attendance_students,
    )


with app.app_context():
    init_db()


if __name__ == "__main__":
    with app.app_context():
        init_db()
    app.run(host="0.0.0.0", port=5000, debug=True)
